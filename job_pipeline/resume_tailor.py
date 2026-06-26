from __future__ import annotations

import re
import shutil
import textwrap

from pathlib import Path
from typing import Any

from .config_loader import load_path_with_fallback
from .docx_qa import convert_docx_to_pdf, find_soffice
from .utils import CONFIG_DIR, PROJECT_ROOT, RESUMES_DIR, TEMPLATES_DIR, flatten_text, list_to_cell, load_yaml, normalize_text_escapes, slugify, today_yyyymmdd



EVIDENCE_BUCKETS = [
    "evidence_backed",
    "project_adjacent",
    "coursework_exposure",
    "learning_backlog",
    "do_not_use",
]


def _norm_keyword(value: Any) -> str:
    return str(value or "").strip().lower()


def _keyword_evidence(keyword_info: dict[str, Any]) -> dict[str, set[str]]:
    raw = keyword_info.get("keyword_evidence") or keyword_info.get("evidence") or {}
    evidence = {bucket: set() for bucket in EVIDENCE_BUCKETS}
    if isinstance(raw, dict):
        for bucket in EVIDENCE_BUCKETS:
            values = raw.get(bucket) or []
            if isinstance(values, str):
                values = [values]
            evidence[bucket] = {_norm_keyword(item) for item in values if _norm_keyword(item)}
    unsafe = evidence["learning_backlog"] | evidence["do_not_use"]
    for bucket in ["evidence_backed", "project_adjacent", "coursework_exposure"]:
        evidence[bucket] = {item for item in evidence[bucket] if item not in unsafe}
    return evidence


def _has_explicit_evidence(keyword_info: dict[str, Any]) -> bool:
    evidence = _keyword_evidence(keyword_info)
    return any(evidence[bucket] for bucket in EVIDENCE_BUCKETS)


def _forbidden_keywords(keyword_info: dict[str, Any]) -> set[str]:
    evidence = _keyword_evidence(keyword_info)
    return evidence["learning_backlog"] | evidence["do_not_use"]


def _coursework_keywords(keyword_info: dict[str, Any]) -> set[str]:
    return _keyword_evidence(keyword_info)["coursework_exposure"]


def _contains_keyword(text: str, keyword: str) -> bool:
    return bool(keyword and keyword in text.lower())


def _contains_any_keyword(text: str, keywords: set[str]) -> bool:
    lower = text.lower()
    return any(keyword and keyword in lower for keyword in keywords)

ROLE_PROFILES = {
    "data": ["python", "sql", "pandas", "data cleaning", "reporting", "dashboard", "market data"],
    "trading_ops": ["market data", "trading system", "reconciliation", "risk control", "reporting", "automation"],
    "risk": ["risk control", "credit risk", "operational risk", "reporting", "data analysis", "financial markets"],
    "capital_markets": ["capital markets", "business analysis", "financial markets", "reporting", "excel"],
    "api": ["api integration", "rest api", "technical operations", "automation", "json", "data pipeline"],
    "crypto_ops": ["crypto market data", "order book", "exchange", "digital assets", "market data", "automation"],
    "general": ["python", "sql", "excel", "market data", "automation", "reporting"],
}


RESUME_PROFILE_PATHS_CONFIG = CONFIG_DIR / "resume_profile_paths.local.yaml"
HUMAN_DOCX_RENDERER = "docx_template"
PDF_RENDERER_DOCX = "docx_to_pdf"
PDF_RENDERER_MARKDOWN_FALLBACK = "markdown_fallback"

PROFILE_ROLE_CATEGORIES = {
    "general_data": "data",
    "business_operations": "general",
    "sales_operations": "general",
    "technical_support": "api",
    "finance_operations": "risk",
    "operations": "trading_ops",
}

SKILL_GROUP_LABELS = {
    "programming_data": "Programming & Data",
    "data_tools": "Programming & Data",
    "technical_tools": "Programming & Data",
    "data_reporting": "Programming & Data",
    "workflow_tools": "AI & Workflow Tools",
    "tools": "AI & Workflow Tools",
    "api_operations": "Platforms & Tools",
    "support_workflows": "Platforms & Tools",
    "market_data": "Market / Risk / Operations Focus",
    "market_analysis": "Market / Risk / Operations Focus",
    "risk_operations": "Market / Risk / Operations Focus",
    "trading_operations": "Market / Risk / Operations Focus",
    "capital_markets": "Market / Risk / Operations Focus",
    "operations": "Market / Risk / Operations Focus",
    "business_operations": "Market / Risk / Operations Focus",
    "business_analysis": "Market / Risk / Operations Focus",
}

SKILL_GROUP_ORDER = [
    "Programming & Data",
    "AI & Workflow Tools",
    "Market / Risk / Operations Focus",
    "Platforms & Tools",
]


def load_master_resume(path: Path) -> dict[str, Any]:
    data = load_path_with_fallback(path)
    return data if isinstance(data, dict) else {}


def _keyword_set(keyword_info: dict[str, Any], role_category: str, *, direct_only: bool = False) -> set[str]:
    evidence = _keyword_evidence(keyword_info)
    forbidden = evidence["learning_backlog"] | evidence["do_not_use"]
    if _has_explicit_evidence(keyword_info):
        keywords = set(evidence["evidence_backed"]) | set(evidence["project_adjacent"])
        if not direct_only:
            keywords.update(evidence["coursework_exposure"])
    else:
        keywords = {_norm_keyword(item) for item in keyword_info.get("top_keywords", []) if _norm_keyword(item)}
    keywords.update(_norm_keyword(item) for item in ROLE_PROFILES.get(role_category, ROLE_PROFILES["general"]))
    return {keyword for keyword in keywords if keyword and keyword not in forbidden}


def _score_text(text: str, keywords: set[str]) -> int:
    lower = text.lower()
    return sum(1 for keyword in keywords if keyword and keyword in lower)


def _rank_items(
    items: list[Any],
    keywords: set[str],
    limit: int | None = None,
    *,
    forbidden_keywords: set[str] | None = None,
) -> list[Any]:
    forbidden_keywords = forbidden_keywords or set()
    safe_items = [item for item in items if not _contains_any_keyword(flatten_text(item), forbidden_keywords)]
    ranked = sorted(safe_items, key=lambda item: _score_text(flatten_text(item), keywords), reverse=True)
    return ranked[:limit] if limit else ranked


def _flatten_skills(master: dict[str, Any]) -> list[str]:
    skills: list[str] = []
    raw_skills = master.get("skills", {})
    if isinstance(raw_skills, dict):
        for values in raw_skills.values():
            if isinstance(values, list):
                skills.extend(str(value) for value in values)
    elif isinstance(raw_skills, list):
        skills.extend(str(value) for value in raw_skills)
    return skills



def _dict_items(value: Any) -> list[dict[str, Any]]:
    if not isinstance(value, list):
        return []
    return [item for item in value if isinstance(item, dict)]

def select_skills(master: dict[str, Any], keywords: set[str], limit: int = 24, keyword_info: dict[str, Any] | None = None) -> list[str]:
    keyword_info = keyword_info or {}
    forbidden = _forbidden_keywords(keyword_info)
    coursework = _coursework_keywords(keyword_info) if _has_explicit_evidence(keyword_info) else set()
    skills = _flatten_skills(master)
    ranked = _rank_items(skills, keywords, forbidden_keywords=forbidden | coursework)
    selected: list[str] = []
    for skill in ranked:
        if skill not in selected:
            selected.append(skill)
        if len(selected) >= limit:
            break
    return selected


def build_summary(master: dict[str, Any], keyword_info: dict[str, Any], role_category: str) -> list[str]:
    keywords = _keyword_set(keyword_info, role_category, direct_only=True)
    forbidden = _forbidden_keywords(keyword_info)
    base = master.get("summary", [])
    ranked = _rank_items(base, keywords, limit=2, forbidden_keywords=forbidden)
    matched_skills = [
        skill
        for skill in select_skills(master, keywords, limit=10, keyword_info=keyword_info)
        if skill.lower() in " ".join(keyword_info.get("top_keywords", [])).lower()
        or _score_text(skill, keywords) > 0
    ][:8]
    summary = list(ranked)
    if matched_skills:
        summary.append("Targeted fit: " + ", ".join(matched_skills) + ".")
    coursework = sorted(_coursework_keywords(keyword_info))
    if coursework and len(summary) < 3:
        summary.append("Coursework exposure to " + ", ".join(coursework[:4]) + ".")
    return summary[:3]


def _format_contact(master: dict[str, Any]) -> str:
    contacts = master.get("contacts", {})
    parts = [
        master.get("location"),
        contacts.get("phone"),
        contacts.get("email"),
        contacts.get("linkedin"),
    ]
    return " | ".join(str(part) for part in parts if part)


def render_markdown_resume(
    master: dict[str, Any],
    job: dict[str, Any],
    keyword_info: dict[str, Any],
) -> str:
    role_category = str(job.get("role_category") or "general")
    keywords = _keyword_set(keyword_info, role_category, direct_only=True)
    forbidden = _forbidden_keywords(keyword_info)
    selected_skills = select_skills(master, keywords, keyword_info=keyword_info)
    projects = _rank_items(_dict_items(master.get("projects")), keywords, limit=3)
    experience = _rank_items(_dict_items(master.get("experience")), keywords, limit=3)

    lines: list[str] = []
    lines.append(f"# {master.get('name', '')}")
    contact = _format_contact(master)
    if contact:
        lines.append(contact)
    if master.get("headline"):
        lines.append("")
        lines.append(f"**{master['headline']}**")
    lines.append("")
    lines.append("## Professional Summary")
    for bullet in build_summary(master, keyword_info, role_category):
        lines.append(f"- {bullet}")

    lines.append("")
    lines.append("## Targeted Skills")
    lines.append(", ".join(selected_skills))

    lines.append("")
    lines.append("## Project Experience")
    for project in projects:
        title = project.get("name", "")
        meta = " | ".join(str(x) for x in [project.get("type"), project.get("dates")] if x)
        lines.append(f"### {title}" + (f" | {meta}" if meta else ""))
        for bullet in _rank_items(project.get("bullets", []), keywords, limit=5, forbidden_keywords=forbidden):
            lines.append(f"- {bullet}")

    lines.append("")
    lines.append("## Professional Experience")
    for role in experience:
        header = " | ".join(str(x) for x in [role.get("title"), role.get("company"), role.get("location"), role.get("dates")] if x)
        lines.append(f"### {header}")
        for bullet in _rank_items(role.get("bullets", []), keywords, limit=4, forbidden_keywords=forbidden):
            lines.append(f"- {bullet}")

    education = _dict_items(master.get("education"))
    if education:
        lines.append("")
        lines.append("## Education")
        for item in education:
            lines.append(f"- {item.get('degree', '')} | {item.get('school', '')} | {item.get('date', '')}")
            if item.get("details"):
                lines.append(f"  - {list_to_cell(item.get('details'))}")

    languages = master.get("languages", [])
    if languages:
        lines.append("")
        lines.append("## Languages")
        for item in languages:
            lines.append(f"- {item}")

    lines.append("")
    lines.append("<!-- Compliance note: generated only from configured resume source files; review manually before applying. -->")
    return "\n".join(lines).strip() + "\n"


def _clear_doc_body(doc: Any) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def _clean_resume_text(value: Any) -> str:
    text = normalize_text_escapes(str(value or ""))
    text = text.replace("\ufffd", "")
    text = text.replace("linkedin.comin", "linkedin.com/in")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _safe_items(values: Any) -> list[str]:
    if isinstance(values, list):
        return [_clean_resume_text(value) for value in values if _clean_resume_text(value)]
    if values:
        return [_clean_resume_text(values)]
    return []


def build_human_summary(master: dict[str, Any], keyword_info: dict[str, Any], role_category: str, limit: int = 3) -> list[str]:
    summary = [
        _clean_resume_text(item)
        for item in build_summary(master, keyword_info, role_category)
        if not _clean_resume_text(item).lower().startswith("targeted fit:")
    ]
    for item in _safe_items(master.get("summary")):
        if item not in summary and not item.lower().startswith("targeted fit:"):
            summary.append(item)
        if len(summary) >= limit:
            break
    return summary[:limit]


def grouped_human_skills(
    master: dict[str, Any],
    keyword_info: dict[str, Any],
    *,
    max_items_per_group: int = 9,
) -> list[tuple[str, list[str]]]:
    forbidden = _forbidden_keywords(keyword_info)
    coursework = _coursework_keywords(keyword_info) if _has_explicit_evidence(keyword_info) else set()
    groups: dict[str, list[str]] = {label: [] for label in SKILL_GROUP_ORDER}
    extra_groups: dict[str, list[str]] = {}
    raw_groups = master.get("skills") or {}
    if isinstance(raw_groups, dict):
        for key, values in raw_groups.items():
            label = SKILL_GROUP_LABELS.get(str(key), str(key).replace("_", " ").title())
            bucket = groups.setdefault(label, []) if label in SKILL_GROUP_ORDER else extra_groups.setdefault(label, [])
            for value in _safe_items(values):
                normalized = _norm_keyword(value)
                if normalized in forbidden or normalized in coursework:
                    continue
                if value not in bucket and len(bucket) < max_items_per_group:
                    bucket.append(value)
    ordered: list[tuple[str, list[str]]] = []
    for label in SKILL_GROUP_ORDER:
        values = groups.get(label) or []
        if values:
            ordered.append((label, values))
    for label, values in extra_groups.items():
        if values:
            ordered.append((label, values[:max_items_per_group]))
    if not ordered:
        fallback = select_skills(master, _keyword_set(keyword_info, "general"), limit=max_items_per_group, keyword_info=keyword_info)
        if fallback:
            ordered.append(("Programming & Data", [_clean_resume_text(item) for item in fallback]))
    return ordered


def _docx_imports() -> dict[str, Any]:
    from docx import Document  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.shared import Inches, Pt  # type: ignore

    return {
        "Document": Document,
        "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
        "WD_TAB_ALIGNMENT": WD_TAB_ALIGNMENT,
        "OxmlElement": OxmlElement,
        "qn": qn,
        "Inches": Inches,
        "Pt": Pt,
    }


def _set_run(run: Any, *, size: float = 9.0, bold: bool = False, italic: bool = False, font: str = "Calibri") -> None:
    imports = _docx_imports()
    run.font.name = font
    run.font.size = imports["Pt"](size)
    run.bold = bold
    run.italic = italic


def _set_paragraph_spacing(paragraph: Any, *, before: float = 0, after: float = 1.5, line_spacing: float = 1.0) -> None:
    imports = _docx_imports()
    paragraph.paragraph_format.space_before = imports["Pt"](before)
    paragraph.paragraph_format.space_after = imports["Pt"](after)
    paragraph.paragraph_format.line_spacing = line_spacing


def _add_bottom_border(paragraph: Any, *, color: str = "666666") -> None:
    imports = _docx_imports()
    OxmlElement = imports["OxmlElement"]
    qn = imports["qn"]
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def _content_width(section: Any) -> Any:
    return section.page_width - section.left_margin - section.right_margin


def _prepare_human_doc(template_path: Path | None = None) -> Any:
    imports = _docx_imports()
    Document = imports["Document"]
    Inches = imports["Inches"]
    Pt = imports["Pt"]
    use_template = bool(template_path and template_path.exists())
    doc = Document(str(template_path)) if use_template else Document()
    if use_template:
        _clear_doc_body(doc)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(9)
    for section in doc.sections:
        section.top_margin = Inches(0.45)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)
    return doc


def _add_centered_line(doc: Any, text: str, *, size: float, bold: bool = False, italic: bool = False, after: float = 1.0) -> None:
    imports = _docx_imports()
    paragraph = doc.add_paragraph()
    paragraph.alignment = imports["WD_ALIGN_PARAGRAPH"].CENTER
    _set_paragraph_spacing(paragraph, after=after)
    run = paragraph.add_run(_clean_resume_text(text))
    _set_run(run, size=size, bold=bold, italic=italic)


def _add_section_heading(doc: Any, title: str) -> None:
    paragraph = doc.add_paragraph()
    _set_paragraph_spacing(paragraph, before=4, after=2)
    run = paragraph.add_run(_clean_resume_text(title).upper())
    _set_run(run, size=9.2, bold=True)
    _add_bottom_border(paragraph)


def _add_bullet(doc: Any, text: str) -> None:
    imports = _docx_imports()
    paragraph = doc.add_paragraph()
    _set_paragraph_spacing(paragraph, after=0.8)
    paragraph.paragraph_format.left_indent = imports["Inches"](0.18)
    paragraph.paragraph_format.first_line_indent = imports["Inches"](-0.18)
    run = paragraph.add_run("Ã¢â‚¬Â¢ ")
    _set_run(run, size=8.6)
    run = paragraph.add_run(_clean_resume_text(text))
    _set_run(run, size=8.6)


def _add_key_value_line(doc: Any, label: str, values: list[str]) -> None:
    paragraph = doc.add_paragraph()
    _set_paragraph_spacing(paragraph, after=0.8)
    label_run = paragraph.add_run(_clean_resume_text(label) + ": ")
    _set_run(label_run, size=8.7, bold=True)
    value_run = paragraph.add_run("; ".join(_clean_resume_text(value) for value in values if _clean_resume_text(value)))
    _set_run(value_run, size=8.7)


def _add_split_header(doc: Any, left: str, right: str = "") -> None:
    imports = _docx_imports()
    paragraph = doc.add_paragraph()
    _set_paragraph_spacing(paragraph, before=1.5, after=0.8)
    paragraph.paragraph_format.tab_stops.add_tab_stop(_content_width(doc.sections[0]), imports["WD_TAB_ALIGNMENT"].RIGHT)
    left_run = paragraph.add_run(_clean_resume_text(left))
    _set_run(left_run, size=8.9, bold=True)
    if right:
        right_run = paragraph.add_run("\t" + _clean_resume_text(right))
        _set_run(right_run, size=8.7, italic=True)


def _human_resume_sections(master: dict[str, Any], job: dict[str, Any], keyword_info: dict[str, Any]) -> dict[str, Any]:
    role_category = str(job.get("role_category") or "general")
    keywords = _keyword_set(keyword_info, role_category, direct_only=True)
    forbidden = _forbidden_keywords(keyword_info)
    projects = _rank_items(_dict_items(master.get("projects")), keywords, limit=3, forbidden_keywords=forbidden)
    experience = _rank_items(_dict_items(master.get("experience")), keywords, limit=3, forbidden_keywords=forbidden)
    return {
        "summary": build_human_summary(master, keyword_info, role_category, limit=3),
        "skills": grouped_human_skills(master, keyword_info),
        "projects": projects,
        "experience": experience,
        "education": _dict_items(master.get("education")),
        "languages": _safe_items(master.get("languages")),
        "keywords": keywords,
        "forbidden": forbidden,
    }


def write_human_docx_resume(
    master: dict[str, Any],
    job: dict[str, Any],
    keyword_info: dict[str, Any],
    out_path: Path,
    *,
    template_path: Path | None = None,
) -> bool:
    try:
        _docx_imports()
    except ModuleNotFoundError:
        return False

    doc = _prepare_human_doc(template_path)
    sections = _human_resume_sections(master, job, keyword_info)
    _add_centered_line(doc, str(master.get("name") or ""), size=15.5, bold=True, after=0.4)
    contact = _format_contact(master)
    if contact:
        _add_centered_line(doc, contact, size=8.4, after=0.2)
    if master.get("headline"):
        _add_centered_line(doc, str(master.get("headline")), size=9.2, italic=True, after=3.0)

    if sections["summary"]:
        _add_section_heading(doc, "Professional Summary")
        for item in sections["summary"][:3]:
            _add_bullet(doc, item)

    if sections["skills"]:
        _add_section_heading(doc, "Skills")
        for label, values in sections["skills"]:
            _add_key_value_line(doc, label, values)

    if sections["projects"]:
        _add_section_heading(doc, "Project Experience")
        for project in sections["projects"]:
            left = " | ".join(_clean_resume_text(value) for value in [project.get("name"), project.get("type")] if _clean_resume_text(value))
            _add_split_header(doc, left, str(project.get("dates") or ""))
            for bullet in _rank_items(project.get("bullets", []), sections["keywords"], limit=5, forbidden_keywords=sections["forbidden"]):
                _add_bullet(doc, str(bullet))

    if sections["experience"]:
        _add_section_heading(doc, "Professional Experience")
        for role in sections["experience"]:
            left = " | ".join(_clean_resume_text(value) for value in [role.get("title"), role.get("company"), role.get("location")] if _clean_resume_text(value))
            _add_split_header(doc, left, str(role.get("dates") or ""))
            for bullet in _rank_items(role.get("bullets", []), sections["keywords"], limit=5, forbidden_keywords=sections["forbidden"]):
                _add_bullet(doc, str(bullet))

    if sections["education"]:
        _add_section_heading(doc, "Education")
        for item in sections["education"]:
            left = " | ".join(_clean_resume_text(value) for value in [item.get("degree"), item.get("school")] if _clean_resume_text(value))
            _add_split_header(doc, left, str(item.get("date") or ""))
            details = _safe_items(item.get("details"))
            if details:
                _add_key_value_line(doc, "Relevant Coursework", details[:6])

    if sections["languages"]:
        _add_section_heading(doc, "Languages")
        _add_key_value_line(doc, "Languages", sections["languages"])

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return True


def write_pdf_from_docx(
    docx_path: Path,
    pdf_path: Path,
    *,
    fallback_markdown: str | None = None,
    allow_fallback: bool = True,
) -> str:
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        if pdf_path.exists():
            pdf_path.unlink()
        created = convert_docx_to_pdf(docx_path, pdf_path.parent, find_soffice())
        if created.resolve() != pdf_path.resolve():
            if pdf_path.exists():
                pdf_path.unlink()
            shutil.move(str(created), str(pdf_path))
        return PDF_RENDERER_DOCX if pdf_path.exists() else ""
    except Exception:
        if allow_fallback and fallback_markdown and write_pdf_from_markdown(fallback_markdown, pdf_path):
            return PDF_RENDERER_MARKDOWN_FALLBACK
        return ""


def _resolve_project_path(path_text: str) -> Path:
    path = Path(path_text)
    return path if path.is_absolute() else PROJECT_ROOT / path


def _profile_keyword_info(master: dict[str, Any]) -> dict[str, Any]:
    keywords = _flatten_skills(master)
    return {"top_keywords": keywords[:30]}


def _profile_job(profile: str, master: dict[str, Any]) -> dict[str, Any]:
    return {
        "canonical_job_id": profile,
        "company": "profile",
        "title": master.get("headline") or profile.replace("_", " ").title(),
        "score": 0,
        "role_category": PROFILE_ROLE_CATEGORIES.get(profile, "general"),
    }


def generate_profile_resumes(
    *,
    profile_config_path: Path = RESUME_PROFILE_PATHS_CONFIG,
    output_dir: Path | None = None,
    make_pdf: bool = True,
    template_path: Path | None = None,
) -> list[dict[str, Any]]:
    config = load_path_with_fallback(profile_config_path)
    profiles = config.get("profiles") if isinstance(config.get("profiles"), dict) else {}
    results: list[dict[str, Any]] = []
    for profile, payload in profiles.items():
        if not isinstance(payload, dict):
            continue
        source_path = _resolve_project_path(str(payload.get("source") or payload.get("yaml") or ""))
        if output_dir is None:
            docx_path = _resolve_project_path(str(payload.get("docx") or f"data/resumes/profiles/{profile}.docx"))
            pdf_path = _resolve_project_path(str(payload.get("pdf") or f"data/resumes/profiles/{profile}.pdf"))
        else:
            docx_path = output_dir / f"{profile}.docx"
            pdf_path = output_dir / f"{profile}.pdf"
        master = load_master_resume(source_path)
        job = _profile_job(str(profile), master)
        keyword_info = _profile_keyword_info(master)
        markdown = render_markdown_resume(master, job, keyword_info)
        docx_created = write_human_docx_resume(
            master,
            job,
            keyword_info,
            docx_path,
            template_path=template_path,
        )
        pdf_renderer = ""
        if make_pdf and docx_created:
            pdf_renderer = write_pdf_from_docx(docx_path, pdf_path, fallback_markdown=markdown)
        elif make_pdf and write_pdf_from_markdown(markdown, pdf_path):
            pdf_renderer = PDF_RENDERER_MARKDOWN_FALLBACK
        results.append(
            {
                "profile": str(profile),
                "source": str(source_path),
                "docx": str(docx_path) if docx_created else "",
                "pdf": str(pdf_path) if pdf_path.exists() else "",
                "docx_renderer": HUMAN_DOCX_RENDERER if docx_created else "",
                "pdf_renderer": pdf_renderer,
            }
        )
    return results


def write_docx_from_markdown(markdown_text: str, out_path: Path, template_path: Path | None = None) -> bool:
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt  # type: ignore
    except ModuleNotFoundError:
        return False

    use_template = bool(template_path and template_path.exists())
    doc = Document(str(template_path)) if use_template else Document()
    if use_template:
        _clear_doc_body(doc)
    else:
        styles = doc.styles
        styles["Normal"].font.name = "Arial"
        styles["Normal"].font.size = Pt(10)
    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line or line.startswith("<!--"):
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("**") and line.endswith("**"):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line.strip("*"))
            run.bold = True
        else:
            doc.add_paragraph(line)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return True


def _wrapped_lines(text: str, *, width: int) -> list[str]:
    if not text:
        return [""]
    return textwrap.wrap(text, width=width, break_long_words=False, replace_whitespace=False) or [""]


def write_pdf_from_markdown(markdown_text: str, out_path: Path) -> bool:
    try:
        import fitz  # type: ignore
    except ModuleNotFoundError:
        return False

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = fitz.open()
    page_width = 612
    page_height = 792
    margin_x = 54
    margin_y = 48
    bottom = page_height - margin_y
    page = doc.new_page(width=page_width, height=page_height)
    y = margin_y

    def new_page() -> None:
        nonlocal page, y
        page = doc.new_page(width=page_width, height=page_height)
        y = margin_y

    def add_line(text: str, *, size: int = 10, indent: int = 0, gap: float = 2.0) -> None:
        nonlocal y
        line_height = size * 1.35
        if y + line_height > bottom:
            new_page()
        page.insert_text((margin_x + indent, y), text, fontsize=size, fontname="helv", color=(0, 0, 0))
        y += line_height + gap

    def add_wrapped(text: str, *, size: int = 10, indent: int = 0, width: int = 92, gap: float = 2.0) -> None:
        for line in _wrapped_lines(text, width=width):
            add_line(line, size=size, indent=indent, gap=gap)

    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if line.startswith("<!--"):
            continue
        if not line:
            y += 5
            if y > bottom:
                new_page()
            continue
        if line.startswith("# "):
            add_wrapped(line[2:], size=17, width=58, gap=5)
        elif line.startswith("## "):
            y += 4
            add_wrapped(line[3:].upper(), size=12, width=68, gap=4)
        elif line.startswith("### "):
            add_wrapped(line[4:], size=10, width=86, gap=3)
        elif line.startswith("- "):
            add_wrapped("- " + line[2:], size=9, indent=12, width=92, gap=1.5)
        elif line.startswith("**") and line.endswith("**"):
            add_wrapped(line.strip("*"), size=10, width=90, gap=3)
        else:
            add_wrapped(line, size=9, width=96, gap=2)

    doc.save(out_path)
    doc.close()
    return True


def generate_resume(
    *,
    master_resume_path: Path,
    job: dict[str, Any],
    keyword_info: dict[str, Any],
    output_dir: Path = RESUMES_DIR,
    output_date: str | None = None,
    make_docx: bool = True,
    make_pdf: bool = True,
    template_path: Path | None = None,
) -> dict[str, Any]:
    master = load_master_resume(master_resume_path)
    date_part = output_date or today_yyyymmdd()
    company = slugify(job.get("company"), 60)
    role = slugify(job.get("title") or job.get("role_category"), 55)
    score = int(job.get("score") or 0)
    resume_id = slugify(job.get("canonical_job_id") or job.get("source_job_id") or job.get("job_id"), 12)
    suffix = f"_{resume_id}" if resume_id != "item" else ""
    stem = f"{role}_{score}{suffix}"
    target_dir = output_dir / company / date_part
    md_path = target_dir / f"{stem}.md"
    docx_path = target_dir / f"{stem}.docx"
    pdf_path = target_dir / f"{stem}.pdf"

    markdown = render_markdown_resume(master, job, keyword_info)
    target_dir.mkdir(parents=True, exist_ok=True)
    md_path.write_text(markdown, encoding="utf-8")
    docx_created = False
    docx_renderer = ""
    if make_docx:
        docx_created = write_human_docx_resume(
            master,
            job,
            keyword_info,
            docx_path,
            template_path=template_path,
        )
        if docx_created:
            docx_renderer = HUMAN_DOCX_RENDERER
        else:
            docx_created = write_docx_from_markdown(
                markdown,
                docx_path,
                template_path=template_path,
            )
            docx_renderer = PDF_RENDERER_MARKDOWN_FALLBACK if docx_created else ""
    pdf_renderer = ""
    if make_pdf:
        if docx_created:
            pdf_renderer = write_pdf_from_docx(docx_path, pdf_path, fallback_markdown=markdown)
        elif write_pdf_from_markdown(markdown, pdf_path):
            pdf_renderer = PDF_RENDERER_MARKDOWN_FALLBACK
    return {
        "directory": str(target_dir),
        "markdown": str(md_path),
        "docx": str(docx_path) if docx_created else "",
        "pdf": str(pdf_path) if pdf_path.exists() else "",
        "docx_renderer": docx_renderer,
        "pdf_renderer": pdf_renderer,
    }
